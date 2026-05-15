
import docx

def extract_everything(file_path):
    doc = docx.Document(file_path)
    content = []
    
    def add_section(items):
        for item in items:
            if hasattr(item, 'text') and item.text.strip():
                content.append(item.text)
            elif hasattr(item, 'paragraphs'):
                for p in item.paragraphs:
                    if p.text.strip():
                        content.append(p.text)

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            content.append(p.text)
            
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        content.append(p.text)
                        
    # Sections (Headers/Footers)
    for section in doc.sections:
        add_section(section.header.paragraphs)
        add_section(section.footer.paragraphs)
        
    return "\n".join(content)

if __name__ == "__main__":
    text = extract_everything("AVCB PALAVRA CHAVE.docx")
    with open("everything_text.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Done")
