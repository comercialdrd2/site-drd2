
import docx

def extract_full_content(file_path):
    doc = docx.Document(file_path)
    content = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                content.append(" | ".join(row_text))
                
    return "\n".join(content)

if __name__ == "__main__":
    text = extract_full_content("AVCB PALAVRA CHAVE.docx")
    with open("full_original_text.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("Done")
